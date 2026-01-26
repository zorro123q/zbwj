import os
import shutil
import unittest
import uuid
from pathlib import Path

from app import create_app
from app.extensions import db
from app.models import File, KbBlock, KbDocument

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


class KbIngestTestCase(unittest.TestCase):
    def setUp(self):
        os.environ["FLASK_ENV"] = "testing"
        self.app = create_app("testing")
        self.app.config["TESTING"] = True

        with self.app.app_context():
            db.create_all()

        self.client = self.app.test_client()
        self.repo_root = Path(self.app.root_path).parent
        self.storage_dir = self.repo_root / "storage"

    def tearDown(self):
        with self.app.app_context():
            db.drop_all()

        if self.storage_dir.exists():
            shutil.rmtree(self.storage_dir)

    @unittest.skipIf(Document is None, "python-docx is required for kb ingest tests")
    def test_ingest_list_delete_kb_docx(self):
        file_id = str(uuid.uuid4())
        rel_path = f"storage/uploads/{file_id}/original.docx"
        abs_path = self.repo_root / rel_path
        abs_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("第一章 总则", level=1)
        doc.add_paragraph("本章内容 1")
        doc.add_heading("1.1 术语", level=2)
        doc.add_paragraph("术语内容 1")
        doc.add_heading("第二章 采购范围", level=1)
        doc.add_paragraph("范围内容 1")
        doc.save(str(abs_path))

        with self.app.app_context():
            rec = File(
                id=file_id,
                filename="demo.docx",
                ext="docx",
                size=abs_path.stat().st_size,
                storage_path=rel_path,
            )
            db.session.add(rec)
            db.session.commit()

        resp = self.client.post("/api/v1/kb/ingest", json={"file_id": file_id, "title": "投标文件"})
        self.assertEqual(resp.status_code, 200)
        data = resp.get_json()
        self.assertIn("doc_id", data)
        self.assertGreaterEqual(data["block_count"], 2)

        doc_id = data["doc_id"]

        with self.app.app_context():
            blocks = KbBlock.query.filter_by(doc_id=doc_id).all()
            self.assertGreaterEqual(len(blocks), 2)
            for block in blocks:
                block_path = self.repo_root / block.block_docx_path
                self.assertTrue(block_path.exists())

            blocks[0].tag = "product_intro"
            db.session.commit()

        list_resp = self.client.get("/api/v1/kb/docs?page=1&page_size=10")
        self.assertEqual(list_resp.status_code, 200)
        list_data = list_resp.get_json()
        self.assertGreaterEqual(list_data["total"], 1)

        search_resp = self.client.post(
            "/api/v1/kb/search",
            json={
                "query": "内容",
                "top_k": 5,
                "by_tag": "product_intro",
                "title_keywords": ["总则", "采购范围"],
            },
        )
        self.assertEqual(search_resp.status_code, 200)
        search_data = search_resp.get_json()
        self.assertGreaterEqual(search_data["total"], 1)

        del_resp = self.client.delete(f"/api/v1/kb/docs/{doc_id}")
        self.assertEqual(del_resp.status_code, 200)

        with self.app.app_context():
            self.assertIsNone(db.session.get(KbDocument, doc_id))
