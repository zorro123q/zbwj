import importlib.util
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from flask import current_app

from domain.kb.retriever import search_blocks
from domain.templates.registry import TemplateRegistry


class WordExportError(ValueError):
    pass


def _get_repo_root() -> Path:
    return Path(current_app.root_path).parent


def _ensure_python_docx():
    if importlib.util.find_spec("docx") is None:
        raise WordExportError("python-docx is required to export word documents")


def _ensure_docxcompose_available() -> bool:
    return importlib.util.find_spec("docxcompose") is not None


def _load_docx(path: Path):
    from docx import Document

    return Document(str(path))


def _merge_block_docx(composer, block_path: Path) -> bool:
    try:
        block_doc = _load_docx(block_path)
        composer.append(block_doc)
        return True
    except Exception:
        return False


def _append_text_fallback(doc, content_text: str) -> None:
    for line in (content_text or "").splitlines():
        if line.strip():
            doc.add_paragraph(line)


def export_by_template(
    *,
    job_id: str,
    template_id: str,
    version: str,
    company_id: Optional[str] = None,
) -> Dict[str, str]:
    job_id = (job_id or "").strip()
    template_id = (template_id or "").strip()
    version = (version or "").strip()
    if not job_id:
        raise WordExportError("job_id is required")
    if not template_id:
        raise WordExportError("template_id is required")
    if not version:
        raise WordExportError("version is required")

    _ensure_python_docx()

    template = TemplateRegistry.get(template_id, version)
    sections = template.get("sections") or []
    if not isinstance(sections, list) or len(sections) == 0:
        raise WordExportError("template sections must be a non-empty array")

    from docx import Document

    output_doc = Document()
    use_composer = _ensure_docxcompose_available()
    composer = None
    if use_composer:
        from docxcompose.composer import Composer

        composer = Composer(output_doc)

    repo_root = _get_repo_root()

    for section in sections:
        title = (section.get("title") or "").strip() if isinstance(section, dict) else ""
        if title:
            output_doc.add_heading(title, level=1)

        pick = section.get("pick") if isinstance(section, dict) else {}
        pick = pick if isinstance(pick, dict) else {}
        by_tag = pick.get("by_tag")
        title_keywords = pick.get("fallback_title_keywords")
        top_k = pick.get("top_k")

        tag_list: List[Optional[str]] = []
        if isinstance(by_tag, str) and by_tag.strip():
            tag_list = [by_tag.strip()]
        elif isinstance(by_tag, list):
            tag_list = [t.strip() for t in by_tag if isinstance(t, str) and t.strip()]
        if not tag_list:
            tag_list = [None]

        merged: Dict[str, Dict[str, Any]] = {}
        for tag in tag_list:
            results = search_blocks(
                query=None,
                top_k=top_k,
                by_tag=tag,
                title_keywords=title_keywords if isinstance(title_keywords, list) else None,
                page=1,
                page_size=top_k or 20,
            )
            for item in results.get("items", []):
                block_id = item.get("block_id")
                if not block_id:
                    continue
                prev = merged.get(block_id)
                if prev is None or int(item.get("score") or 0) > int(prev.get("score") or 0):
                    merged[block_id] = item

        sorted_items = sorted(
            merged.values(), key=lambda x: int(x.get("score") or 0), reverse=True
        )
        if isinstance(top_k, int) and top_k > 0:
            sorted_items = sorted_items[:top_k]

        if not sorted_items:
            raise WordExportError(f"no blocks found for section: {title or 'untitled'}")

        for block in sorted_items:
            rel_path = (block.get("content_docx_path") or "").replace("\\", "/").lstrip("/")
            block_path = (repo_root / Path(rel_path)).resolve()
            if composer and block_path.exists():
                if _merge_block_docx(composer, block_path):
                    continue
            _append_text_fallback(output_doc, block.get("content_text") or "")

    artifacts_dir = f"{current_app.config.get('ARTIFACT_STORAGE_DIR', 'storage/artifacts')}/{job_id}"
    artifacts_dir = artifacts_dir.replace("\\", "/")
    rel_output = f"{artifacts_dir}/tender_response.docx"
    abs_output = repo_root / Path(rel_output)
    abs_output.parent.mkdir(parents=True, exist_ok=True)

    if composer:
        composer.save(str(abs_output))
    else:
        output_doc.save(str(abs_output))

    return {"docx_path": rel_output}
