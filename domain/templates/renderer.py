import os
import uuid
import logging
from typing import List, Dict, Any, Union
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")

logger = logging.getLogger(__name__)


def render_docx_template(data: Union[List[Dict[str, Any]], Dict[str, Any]], template_path: str = None) -> str:
    """
    智能标书与表格渲染引擎
    """
    if template_path and os.path.exists(template_path):
        try:
            doc = Document(template_path)
            logger.info(f"Loaded template: {template_path}")
        except Exception as e:
            logger.warning(f"Failed to load template, creating new doc. Error: {e}")
            doc = Document()
    else:
        doc = Document()

    # 尝试设置全局中文字体支持
    try:
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass

    # 兼容旧版本：如果传进来的是个列表，说明是旧版单纯生成表格的逻辑
    if isinstance(data, list):
        requirements = data
        context_data = {}
        is_full_doc = False
    else:
        # 新版：传入的是字典，包含各类正文段落和需求表
        requirements = data.get("requirements", [])
        context_data = data
        is_full_doc = True

    # ==========================================
    # 模块 A：写入标书标准正文章节
    # ==========================================
    if is_full_doc:
        if not template_path:
            doc.add_heading('投标文件 (系统智能生成)', 0)

        # 1. 商务响应部分
        if context_data.get("company_profile"):
            doc.add_heading('第一部分 商务响应', 1)
            doc.add_heading('1.1 企业简介与综合实力', 2)
            doc.add_paragraph(context_data.get("company_profile"))

        # 2. 技术响应部分
        if context_data.get("tech_solution"):
            doc.add_heading('第二部分 技术方案响应', 1)
            doc.add_heading('2.1 核心技术与系统架构', 2)
            doc.add_paragraph(context_data.get("tech_solution"))

        # 3. 实施与售后部分
        if context_data.get("implementation") or context_data.get("after_sales"):
            doc.add_heading('第三部分 项目实施与售后服务', 1)
            if context_data.get("implementation"):
                doc.add_heading('3.1 实施与培训计划', 2)
                doc.add_paragraph(context_data.get("implementation"))
            if context_data.get("after_sales"):
                doc.add_heading('3.2 售后服务与故障响应', 2)
                doc.add_paragraph(context_data.get("after_sales"))

        # 准备写入表格
        doc.add_heading('第四部分 招标需求点对点响应表', 1)
    else:
        if not template_path:
            doc.add_heading('评审办法索引表', 0)

    # ==========================================
    # 模块 B：写入点对点响应表格
    # ==========================================
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = 'Table Grid'
    except Exception as e:
        logger.warning(f"Style 'Table Grid' not found. Using default. Error: {e}")

    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号/类别'
    hdr_cells[1].text = '招标参数要求'
    hdr_cells[2].text = '我方详细响应方案'
    hdr_cells[3].text = '偏离情况'

    def get_val(item, keys):
        if isinstance(item, dict):
            for k in keys:
                if k in item and item[k]: return str(item[k])
        else:
            for k in keys:
                if hasattr(item, k) and getattr(item, k): return str(getattr(item, k))
        return ''

    # 填充表格内容
    for idx, req in enumerate(requirements):
        row_cells = table.add_row().cells

        cat = get_val(req, ['category', '评审大类', '大类']) or str(idx + 1)
        item = get_val(req, ['item', '评审项', 'desc', 'description', 'content'])

        # 兼容新老字段提取方案
        if "response_text" in req:
            val = req["response_text"]  # 新版标书的详细响应长文
        else:
            val = get_val(req, ['evidence', 'value', '响应内容'])  # 旧版简短索引

        row_cells[0].text = cat
        row_cells[1].text = item
        row_cells[2].text = val
        # 作为正式标书，偏离情况默认自动填写“无偏离”，代表完全满足要求
        row_cells[3].text = '无偏离'

    # ==========================================
    # 模块 C：保存文件
    # ==========================================
    try:
        from flask import current_app
        root = Path(current_app.root_path).parent
    except:
        root = Path(os.getcwd())

    out_dir = root / "storage" / "artifacts" / "review_index"
    out_dir.mkdir(parents=True, exist_ok=True)

    prefix = "bidding_document" if is_full_doc else "review_index"
    filename = f"{prefix}_{uuid.uuid4().hex[:8]}.docx"
    out_path = out_dir / filename

    doc.save(str(out_path))
    logger.info(f"Generated docx at: {out_path}")

    return str(out_path)