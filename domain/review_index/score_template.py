# domain/review_index/score_template.py
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from flask import current_app


@dataclass
class ScoreTemplateRow:
    score_major: str
    score_minor: str
    score_rule: str
    evidence_materials: str
    pages: str


def _repo_root() -> Path:
    return Path(current_app.root_path).parent


def _resolve(p: str) -> Path:
    pp = Path(p)
    if pp.is_absolute():
        return pp
    return (_repo_root() / pp).resolve()


def _cell_text(cell) -> str:
    # 兼容合并单元格：把 cell 里所有段落拼起来
    lines = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    return "\n".join(lines).strip()


def _norm_header(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace(" ", "")
    return s


def _map_header_indices(header_cells: List[str]) -> Dict[str, int]:
    """
    将模板表头映射到逻辑列：
      score_major / score_minor / score_rule / evidence / pages
    """
    header_norm = [_norm_header(x) for x in header_cells]

    candidates = {
        "score_major": ["评分大类", "评分大类(分)", "评分大类（分）", "大类"],
        "score_minor": ["评分小类", "小类"],
        "score_rule": ["评分类别", "评分规则", "评分标准", "类别"],
        "evidence": ["有效证明材料", "证明材料", "材料要求"],
        "pages": ["证明材料页码", "页码", "材料页码"],
    }

    def find_idx(keys: List[str]) -> Optional[int]:
        for k in keys:
            kn = _norm_header(k)
            for i, h in enumerate(header_norm):
                if kn and kn == h:
                    return i
        # 允许“包含匹配”（比如表头写了“评分大类（10分）”）
        for k in keys:
            kn = _norm_header(k)
            for i, h in enumerate(header_norm):
                if kn and (kn in h):
                    return i
        return None

    m: Dict[str, int] = {}
    for logical, keys in candidates.items():
        idx = find_idx(keys)
        if idx is not None:
            m[logical] = idx

    return m


def load_score_template_docx(template_path: str) -> List[ScoreTemplateRow]:
    """
    读取评分模板 docx（第一个表格），按表头映射列，兼容合并单元格。
    """
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required to parse docx") from exc

    abs_path = _resolve(template_path)
    if not abs_path.exists():
        raise FileNotFoundError(f"template docx not found: {abs_path}")

    doc = Document(str(abs_path))
    if not doc.tables:
        raise ValueError("template docx has no table")

    table = doc.tables[0]
    if len(table.rows) < 2:
        return []

    header_cells = [_cell_text(c) for c in table.rows[0].cells]
    col_map = _map_header_indices(header_cells)

    # 最少要能定位到“评分大类”
    if "score_major" not in col_map:
        raise ValueError(f"template header mismatch, cannot find 评分大类. header={header_cells}")

    def get_cell(cells: List[str], key: str) -> str:
        idx = col_map.get(key)
        if idx is None:
            return ""
        return cells[idx] if idx < len(cells) else ""

    rows: List[ScoreTemplateRow] = []
    for r in table.rows[1:]:
        cells = [_cell_text(c) for c in r.cells]

        rows.append(
            ScoreTemplateRow(
                score_major=get_cell(cells, "score_major").strip(),
                score_minor=get_cell(cells, "score_minor").strip(),
                score_rule=get_cell(cells, "score_rule").strip(),
                evidence_materials=get_cell(cells, "evidence").strip(),
                pages=get_cell(cells, "pages").strip(),
            )
        )

    # 过滤掉全空行
    rows = [x for x in rows if (x.score_major or x.score_minor or x.score_rule or x.evidence_materials or x.pages)]
    return rows
