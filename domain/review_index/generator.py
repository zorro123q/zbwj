# domain/review_index/generator.py
from __future__ import annotations

import uuid
from pathlib import Path
from typing import List, Optional, Tuple

from flask import current_app

# 引入新的清洗函数
from domain.review_index.requirements import (
    load_requirements_xlsx,
    clean_and_aggregate_requirements,
    AggregatedReq
)
from domain.review_index.score_template import load_score_template_docx, ScoreTemplateRow
from domain.review_index.kb_evidence import retrieve_evidence_blocks


def _repo_root() -> Path:
    return Path(current_app.root_path).parent


def _resolve(p: str) -> Path:
    pp = Path(p)
    if pp.is_absolute():
        return pp
    return (_repo_root() / pp).resolve()


def _pick_template_req(tpl: ScoreTemplateRow) -> str:
    x = (tpl.evidence_materials or "").strip()
    if x: return x
    x = (tpl.score_rule or "").strip()
    if x: return x
    x = (tpl.score_minor or "").strip()
    if x: return x
    return "-"


def generate_review_index_docx(
        *,
        xlsx_path: str,
        template_docx_path: str,
        kb_tag: Optional[str] = None,
        evidence_top_n: int = 3,
        excerpt_len: int = 800,
) -> Path:
    """
    输出《评审办法索引目录.docx》
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except Exception as exc:
        raise RuntimeError("python-docx is required to write docx") from exc

    # 1. 读原始数据
    raw_reqs = load_requirements_xlsx(xlsx_path=xlsx_path)
    tpl_rows = load_score_template_docx(template_docx_path)

    # 2. 执行清洗与聚合 (New Logic)
    agg_reqs = clean_and_aggregate_requirements(raw_reqs)

    # 尝试提取项目名称 (从聚合结果里找 "基本信息" 或 原始数据找)
    project_name = ""
    for r in raw_reqs:
        if "项目名称" in r.item:
            project_name = r.value
            break

    out_rel = f"storage/kb/exports/review_index_{uuid.uuid4().hex}.docx"
    out_abs = _resolve(out_rel)
    out_abs.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title
    h1 = doc.add_heading("评审办法索引目录", level=1)
    h1.alignment = 1  # Center

    if project_name:
        p = doc.add_paragraph()
        run = p.add_run(f"项目名称：{project_name}")
        run.bold = True
        p.alignment = 1

    doc.add_paragraph(f"生成时间：{uuid.uuid4().hex[:8]}")
    doc.add_paragraph(f"KB tag：{kb_tag or 'ALL'}")
    doc.add_paragraph("-" * 30)

    # =========================
    # 一、招标文件要求摘要 (Refactored Output)
    # =========================
    doc.add_heading("一、招标文件重点要求摘要", level=2)
    doc.add_paragraph("以下内容经自动清洗聚合，页码由行号估算（每页约50行）：")

    if not agg_reqs:
        doc.add_paragraph("（result.xlsx 中未提取到有效要求）")

    for req in agg_reqs:
        # 1. Category 标题
        doc.add_heading(req.category, level=3)

        # 2. 内容摘要 (Content Summary)
        if req.content_summary:
            p = doc.add_paragraph(req.content_summary)
            # 可选：设置一点缩进
            p.paragraph_format.left_indent = Pt(10)
        else:
            doc.add_paragraph("（无具体内容）")

        # 3. 来源 (References)
        if req.references:
            # 格式化展示：[Page:4 3.1 评分标准, Page:6 ...]
            ref_str = "  ".join(req.references)
            ref_p = doc.add_paragraph(f"来源定位：{ref_str}")
            ref_p.paragraph_format.left_indent = Pt(10)
            # 设置灰色字体
            for run in ref_p.runs:
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.font.size = Pt(9)

    doc.add_page_break()

    # =========================
    # 二、评审办法索引目录表 (Keep existing logic mostly)
    # =========================
    doc.add_heading("二、评审办法索引目录（评分模板 + KB证据摘录）", level=2)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "评分大类"
    hdr[1].text = "评分小类"
    hdr[2].text = "评分类别"
    hdr[3].text = "有效证明材料\n(模板要求 + KB内容)"
    hdr[4].text = "定位\n(页码/段落)"

    # 用于附录
    all_evidence: List[Tuple[str, List[dict]]] = []

    for tpl in tpl_rows:
        row_dict = {
            "score_major": tpl.score_major,
            "score_minor": tpl.score_minor,
            "score_rule": tpl.score_rule,
            "evidence_materials": tpl.evidence_materials,
            "pages": tpl.pages,
        }

        hits = retrieve_evidence_blocks(
            score_row=row_dict,
            tag=kb_tag,
            top_n=evidence_top_n,
            excerpt_len=excerpt_len,
        )
        all_evidence.append((tpl.score_major or "（无标题评分大类）", hits))

        tpl_req = _pick_template_req(tpl)

        evidence_lines: List[str] = []
        evidence_lines.append("【模板要求】")
        evidence_lines.append(tpl_req)
        evidence_lines.append("")
        evidence_lines.append("【KB命中】")

        if hits:
            for i, h in enumerate(hits, start=1):
                evidence_lines.append(
                    f"[{i}] {h.get('section_title', '')}"
                )
                excerpt = (h.get("excerpt") or "").strip()
                if excerpt:
                    # 截断一下避免表格太长
                    display_ex = excerpt[:100] + "..." if len(excerpt) > 100 else excerpt
                    evidence_lines.append(f"   {display_ex}")
        else:
            evidence_lines.append("（未命中）")

        # —— 页码/定位列
        page_lines: List[str] = []
        page_lines.append(f"TPL: {tpl.pages or '-'}")
        if hits:
            for i, h in enumerate(hits, start=1):
                # 尝试把 block docx path 转为简单的 ID 或 页码(如果有)
                page_lines.append(f"KB[{i}]")
        else:
            page_lines.append("-")

        rr = table.add_row().cells
        rr[0].text = tpl.score_major or ""
        rr[1].text = tpl.score_minor or ""
        rr[2].text = tpl.score_rule or ""
        rr[3].text = "\n".join(evidence_lines)
        rr[4].text = "\n".join(page_lines)

    # =========================
    # 附录
    # =========================
    doc.add_page_break()
    doc.add_heading("附录：知识库证据详单", level=2)

    for major, hits in all_evidence:
        if not hits: continue
        doc.add_heading(major, level=3)
        for i, h in enumerate(hits, start=1):
            doc.add_paragraph(f"{i}) {h.get('section_title', '')}", style="Heading 4")
            ex = (h.get("excerpt") or "").strip()
            p = doc.add_paragraph(ex if ex else "（空摘录）")
            p.paragraph_format.left_indent = Pt(12)

    doc.save(str(out_abs))
    return out_abs