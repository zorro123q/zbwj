import uuid
from pathlib import Path
from typing import Iterable, Optional, List, Dict, Any

from flask import current_app

from domain.kb.retriever import search_blocks


class KbExportError(RuntimeError):
    pass


def _repo_root() -> Path:
    # 假设当前运行在 app 上下文，向上两级找到根目录
    return Path(current_app.root_path).parent


def export_search_to_docx(
        *,
        query: str,
        top_k: int = 50,
        by_tag: Optional[str] = None,
        title_keywords: Optional[Iterable[str]] = None,
) -> str:
    q = (query or "").strip()
    if not q:
        raise KbExportError("query is required")

    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise KbExportError("python-docx is required to export docx") from exc

    top_k = max(int(top_k or 0), 0)
    if top_k == 0:
        top_k = 50

    # 分页拉取
    items: List[Dict[str, Any]] = []
    page = 1
    page_size = min(100, top_k)

    while len(items) < top_k:
        res = search_blocks(
            query=q,
            top_k=top_k,
            by_tag=by_tag,
            title_keywords=title_keywords,
            page=page,
            page_size=page_size,
        )
        batch = res.get("items") or []
        if not batch:
            break
        items.extend(batch)
        if len(items) >= res.get("total", 0):
            break
        page += 1

    items = items[:top_k]

    # 构造输出路径
    out_rel = f"storage/kb/exports/kb_export_{uuid.uuid4().hex}.docx"
    # 兼容处理：确保目录存在
    out_abs = _repo_root() / Path(out_rel)
    out_abs.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()
    doc.add_heading(f"KB Export: {q}", level=1)
    doc.add_paragraph(f"tag={by_tag}  top_k={top_k}")

    for i, it in enumerate(items, start=1):
        # 使用 filename 作为标题
        title = it.get('filename') or "Unknown File"
        doc.add_heading(f"{i}. {title}", level=2)

        # 输出 file_id 和元数据
        meta_info = f"file_id={it.get('file_id')}  block_id={it.get('block_id')}  score={it.get('score')}"
        if it.get('meta'):
            meta_info += f"\nmeta={it.get('meta')}"

        doc.add_paragraph(meta_info)

        content = it.get("content_text") or ""
        # 简单处理换行
        for line in content.splitlines():
            if line.strip():
                doc.add_paragraph(line)

        doc.add_paragraph("-" * 40)  # 分隔线

    doc.save(str(out_abs))
    return str(out_abs)