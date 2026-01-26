# domain/kb/exporter.py
from __future__ import annotations

import uuid
from pathlib import Path
from typing import Optional

from flask import current_app
from docx import Document as DocxDocument

from app.extensions import db
from app.models import KbBlock, KbDocument


def _exports_dir() -> Path:
    base = Path(current_app.instance_path) / "kb_storage" / "exports"
    base.mkdir(parents=True, exist_ok=True)
    return base


def search_blocks_simple(query: str, top_k: int = 10, by_tag: Optional[str] = None):
    q = db.session.query(KbBlock, KbDocument).join(KbDocument, KbBlock.doc_id == KbDocument.id)

    if by_tag:
        q = q.filter(KbBlock.tag == by_tag)

    like = f"%{query}%"
    q = q.filter(KbBlock.content_text.like(like)).order_by(KbDocument.created_at.desc()).limit(top_k)

    items = []
    for b, d in q.all():
        items.append(
            {
                "doc_id": d.id,
                "title": d.title,
                "block_id": b.id,
                "tag": b.tag,
                "section_title": b.section_title,
                "section_path": b.section_path,
                "content_text": b.content_text,
                "block_docx_path": b.block_docx_path,
            }
        )
    return items


def export_search_to_docx(query: str, top_k: int = 10, by_tag: Optional[str] = None) -> str:
    items = search_blocks_simple(query=query, top_k=top_k, by_tag=by_tag)

    out_path = _exports_dir() / f"kb_export_{uuid.uuid4().hex}.docx"
    doc = DocxDocument()
    doc.add_heading(f"KB Export - query: {query}", level=1)
    doc.add_paragraph(f"top_k={top_k}, by_tag={by_tag}")

    for i, it in enumerate(items, start=1):
        doc.add_heading(f"{i}. {it['title']} | {it['section_title']}", level=2)
        doc.add_paragraph(f"doc_id={it['doc_id']}  block_id={it['block_id']}  tag={it['tag']}")
        for line in (it["content_text"] or "").splitlines():
            doc.add_paragraph(line)

    doc.save(str(out_path))
    return str(out_path)
