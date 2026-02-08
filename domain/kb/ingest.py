# domain/kb/ingest.py
import re
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional

from flask import current_app

from app.extensions import db
from app.models import File, KbBlock, KbDocument


class KbIngestError(ValueError):
    pass


HEADING_NUMBER_RE = re.compile(r"^(\d+(?:\.\d+)*)([\s、.．]+)(.+)")
CN_HEADING_RE = re.compile(r"^([一二三四五六七八九十]+)[、.．](.+)")


@dataclass
class BlockDraft:
    section_title: str
    section_path: str
    content_lines: List[str]
    start_idx: int
    end_idx: int


def _load_docx(path: Path):
    try:
        from docx import Document
    except Exception as exc:
        raise KbIngestError("python-docx is required to parse docx") from exc
    try:
        return Document(str(path))
    except Exception as exc:
        raise KbIngestError("failed to parse docx file") from exc


def _detect_heading_level(paragraph) -> Optional[int]:
    style = getattr(paragraph, "style", None)
    style_name = getattr(style, "name", "") or ""
    if style_name.startswith("Heading "):
        level_str = style_name.replace("Heading ", "").strip()
        if level_str.isdigit():
            return int(level_str)

    text = (paragraph.text or "").strip()
    if not text:
        return None

    match = HEADING_NUMBER_RE.match(text)
    if match:
        segments = match.group(1).split(".")
        return max(1, min(len(segments), 6))

    match = CN_HEADING_RE.match(text)
    if match:
        return 1

    return None


def _normalize_heading_text(text: str) -> str:
    text = (text or "").strip()
    match = HEADING_NUMBER_RE.match(text)
    if match:
        return match.group(3).strip()
    match = CN_HEADING_RE.match(text)
    if match:
        return match.group(2).strip()
    return text


def _iter_blocks(doc) -> List[BlockDraft]:
    blocks: List[BlockDraft] = []
    stack: List[str] = []
    current_block: Optional[BlockDraft] = None

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = (paragraph.text or "").strip()
        if not text:
            continue

        level = _detect_heading_level(paragraph)
        if level is not None:
            if current_block is not None:
                blocks.append(current_block)

            stack = stack[: max(level - 1, 0)]
            heading_text = _normalize_heading_text(text)
            stack.append(heading_text or f"Section {idx}")
            section_path = " / ".join(stack)

            current_block = BlockDraft(
                section_title=heading_text or stack[-1],
                section_path=section_path,
                content_lines=[],
                start_idx=idx,
                end_idx=idx,
            )
            continue

        if current_block is None:
            current_block = BlockDraft(
                section_title="正文",
                section_path="正文",
                content_lines=[],
                start_idx=idx,
                end_idx=idx,
            )

        current_block.content_lines.append(text)
        current_block.end_idx = idx

    if current_block is not None:
        blocks.append(current_block)

    return blocks


def _repo_root() -> Path:
    # 【Fix 2】使用 config 中的 PROJECT_ROOT，避免相对路径计算错误
    root = current_app.config.get("PROJECT_ROOT")
    if root:
        return Path(root)
    # 兜底逻辑：如果配置未加载，则回退到原逻辑
    return Path(current_app.root_path).parent


def _write_block_docx(doc_id: str, block_id: str, title: str, content_lines: Iterable[str]) -> str:
    """
    把每个切片导出成 docx：
      storage/kb/blocks/<doc_id>/<block_id>.docx
    返回写入数据库的相对路径（rel_path）
    """
    try:
        from docx import Document
    except Exception as exc:
        raise KbIngestError("python-docx is required to write docx") from exc

    doc = Document()
    doc.add_heading(title, level=1)
    for line in content_lines:
        doc.add_paragraph(line)

    rel_path = f"storage/kb/blocks/{doc_id}/{block_id}.docx"
    abs_path = _repo_root() / Path(rel_path)
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(abs_path))
    return rel_path


def ingest_kb(file_id: str, title: Optional[str]) -> Dict[str, int]:
    """
    在线模式：依赖 files 表，通过 file_id 找到 stored.storage_path，再切片。
    """
    file_id = (file_id or "").strip()
    if not file_id:
        raise KbIngestError("file_id is required")

    stored = db.session.get(File, file_id)
    if stored is None:
        raise KbIngestError("file_id not found")

    ext = (stored.ext or "").lower().strip()
    if ext != "docx":
        raise KbIngestError("only docx is supported for now")

    # 注意：stored.storage_path 现在建议在存入时就存相对路径
    # 这里使用 _repo_root() 拼接绝对路径
    abs_path = _repo_root() / Path(stored.storage_path)
    if not abs_path.exists():
        raise KbIngestError(f"source file not found on disk: {abs_path}")

    doc = _load_docx(abs_path)
    blocks = _iter_blocks(doc)
    if not blocks:
        raise KbIngestError("no content blocks found in docx")

    doc_id = str(uuid.uuid4())
    doc_title = (title or "").strip() or (stored.filename or "Untitled")
    doc_rec = KbDocument(id=doc_id, file_id=file_id, title=doc_title)
    db.session.add(doc_rec)

    # 关键：先 flush，避免 kb_blocks 外键 1452
    db.session.flush()

    for block in blocks:
        block_id = str(uuid.uuid4())
        block_docx_path = _write_block_docx(doc_id, block_id, block.section_title, block.content_lines)
        content_text = "\n".join(block.content_lines)

        block_rec = KbBlock(
            id=block_id,
            doc_id=doc_id,
            tag=None,
            section_title=block.section_title,
            section_path=block.section_path,
            content_text=content_text,
            start_idx=int(block.start_idx),
            end_idx=int(block.end_idx),
            block_docx_path=block_docx_path,
        )
        db.session.add(block_rec)

    db.session.commit()
    return {"doc_id": doc_id, "block_count": len(blocks)}


def ingest_kb_from_path(file_path: str, title: Optional[str] = None, tag: Optional[str] = None) -> Dict[str, int]:
    """
    离线模式：不需要 file_id，直接读取本地 docx 路径并切片入库。
    """
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        raise KbIngestError(f"source file not found: {p}")
    if p.suffix.lower() != ".docx":
        raise KbIngestError("only docx is supported for now")

    doc = _load_docx(p)
    blocks = _iter_blocks(doc)
    if not blocks:
        raise KbIngestError("no content blocks found in docx")

    doc_id = str(uuid.uuid4())
    pseudo_file_id = str(uuid.uuid4())  # KbDocument.file_id 仍必填，这里用占位 UUID
    doc_title = (title or "").strip() or p.stem

    doc_rec = KbDocument(id=doc_id, file_id=pseudo_file_id, title=doc_title)
    db.session.add(doc_rec)

    # 关键：先 flush，保证 kb_documents 先插入
    db.session.flush()

    for block in blocks:
        block_id = str(uuid.uuid4())
        block_docx_path = _write_block_docx(doc_id, block_id, block.section_title, block.content_lines)
        content_text = "\n".join(block.content_lines)

        block_rec = KbBlock(
            id=block_id,
            doc_id=doc_id,
            tag=(tag or None),
            section_title=block.section_title,
            section_path=block.section_path,
            content_text=content_text,
            start_idx=int(block.start_idx),
            end_idx=int(block.end_idx),
            block_docx_path=block_docx_path,
        )
        db.session.add(block_rec)

    db.session.commit()
    return {"doc_id": doc_id, "block_count": len(blocks)}


def list_docs(page: int, page_size: int) -> Dict[str, object]:
    page = max(int(page), 1)
    page_size = max(min(int(page_size), 100), 1)

    query = KbDocument.query.order_by(KbDocument.created_at.desc())
    total = query.count()
    docs = query.offset((page - 1) * page_size).limit(page_size).all()

    items = []
    for doc in docs:
        block_count = KbBlock.query.filter_by(doc_id=doc.id).count()
        items.append(
            {
                "doc_id": doc.id,
                "file_id": doc.file_id,
                "title": doc.title,
                "block_count": block_count,
                "created_at": doc.created_at.isoformat() if doc.created_at else None,
            }
        )

    return {"page": page, "page_size": page_size, "total": total, "items": items}


def delete_doc(doc_id: str) -> Dict[str, int]:
    doc_id = (doc_id or "").strip()
    if not doc_id:
        raise KbIngestError("doc_id is required")

    doc = db.session.get(KbDocument, doc_id)
    if doc is None:
        raise KbIngestError("doc not found")

    blocks = KbBlock.query.filter_by(doc_id=doc_id).all()

    repo_root = _repo_root()
    for block in blocks:
        rel_path = (block.block_docx_path or "").replace("\\", "/").lstrip("/")
        abs_path = (repo_root / Path(rel_path)).resolve()
        if abs_path.exists() and abs_path.is_file():
            abs_path.unlink()
        db.session.delete(block)

    db.session.delete(doc)
    db.session.commit()

    return {"deleted_blocks": len(blocks)}