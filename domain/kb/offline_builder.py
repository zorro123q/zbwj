# domain/kb/offline_builder.py
from __future__ import annotations

import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Tuple

from flask import current_app
from docx import Document as DocxDocument
from sqlalchemy.exc import SQLAlchemyError

from app.extensions import db
from app.models import KbDocument, KbBlock


class KbOfflineBuildError(Exception):
    pass


def _kb_storage_dir() -> Path:
    base = Path(current_app.instance_path) / "kb_storage"
    (base / "blocks").mkdir(parents=True, exist_ok=True)
    (base / "exports").mkdir(parents=True, exist_ok=True)
    return base


def _read_docx_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _chunk_text(text: str, chunk_chars: int = 1500, overlap: int = 200) -> List[Tuple[int, int, str]]:
    n = len(text)
    if n == 0:
        return []
    chunks: List[Tuple[int, int, str]] = []
    start = 0
    while start < n:
        end = min(n, start + chunk_chars)
        chunks.append((start, end, text[start:end]))
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks


def _write_block_docx(block_path: Path, title: str, content: str) -> None:
    block_path.parent.mkdir(parents=True, exist_ok=True)
    d = DocxDocument()
    d.add_heading(title, level=1)
    for line in content.splitlines():
        d.add_paragraph(line)
    d.save(str(block_path))


def ingest_docx_file_offline(
    file_path: str,
    title: Optional[str] = None,
    tag: Optional[str] = None,
    chunk_chars: int = 1500,
    overlap: int = 200,
) -> dict:
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        raise KbOfflineBuildError(f"file not found: {p}")
    if p.suffix.lower() != ".docx":
        raise KbOfflineBuildError(f"only .docx supported: {p}")

    storage = _kb_storage_dir()
    doc_id = str(uuid.uuid4())
    file_id = str(uuid.uuid4())
    doc_title = title or p.stem

    text = _read_docx_text(p)
    if not text.strip():
        raise KbOfflineBuildError(f"no text extracted: {p}")

    chunks = _chunk_text(text, chunk_chars=chunk_chars, overlap=overlap)
    if not chunks:
        raise KbOfflineBuildError(f"no chunks generated: {p}")

    # 记录生成过哪些 block 文件，失败时可清理
    created_block_files: List[Path] = []

    try:
        # ✅ 先插父表
        kb_doc = KbDocument(
            id=doc_id,
            file_id=file_id,
            title=doc_title,
            created_at=datetime.utcnow(),  # 更稳：即使 DB 没 server_default 也不为 NULL
        )
        db.session.add(kb_doc)

        # ✅ 强制 flush：确保 kb_documents 先落库，外键才能过
        db.session.flush()

        blocks_dir = storage / "blocks" / doc_id
        created = 0

        for idx, (start_i, end_i, chunk) in enumerate(chunks, start=1):
            block_id = str(uuid.uuid4())
            section_title = f"{doc_title} - chunk {idx}"
            section_path = f"/chunk/{idx}"

            block_docx_path = blocks_dir / f"{block_id}.docx"
            _write_block_docx(block_docx_path, section_title, chunk)
            created_block_files.append(block_docx_path)

            kb_block = KbBlock(
                id=block_id,
                doc_id=doc_id,
                tag=tag,
                section_title=section_title,
                section_path=section_path,
                content_text=chunk,
                start_idx=start_i,
                end_idx=end_i,
                block_docx_path=str(block_docx_path),
            )
            db.session.add(kb_block)
            created += 1

        db.session.commit()
        return {
            "status": "ok",
            "doc_id": doc_id,
            "file_id": file_id,
            "title": doc_title,
            "blocks_created": created,
            "source_path": str(p),
        }

    except Exception as e:
        db.session.rollback()
        # 可选：清理刚生成的 block docx，避免污染目录
        for fp in created_block_files:
            try:
                fp.unlink(missing_ok=True)
            except Exception:
                pass
        raise KbOfflineBuildError(str(e)) from e


def ingest_dir_offline(
    root: str,
    pattern: str = "*.docx",
    tag: Optional[str] = None,
    chunk_chars: int = 1500,
    overlap: int = 200,
    exclude_subdirs: Optional[List[str]] = None,
) -> list[dict]:
    base = Path(root).expanduser().resolve()
    if not base.exists():
        raise KbOfflineBuildError(f"root not found: {base}")

    exclude_subdirs = exclude_subdirs or ["instance/kb_storage", "instance\\kb_storage"]
    exclude_paths = [(base / x).resolve() for x in exclude_subdirs]

    results: list[dict] = []

    for p in base.rglob(pattern):
        rp = p.resolve()
        # ✅ 排除 instance/kb_storage 下的 docx（这些是你生成的 block 文件）
        if any(str(rp).startswith(str(ex)) for ex in exclude_paths):
            continue

        try:
            r = ingest_docx_file_offline(
                file_path=str(rp),
                title=rp.stem,
                tag=tag,
                chunk_chars=chunk_chars,
                overlap=overlap,
            )
            results.append(r)
        except Exception as e:
            results.append({"status": "failed", "path": str(rp), "error": str(e)})
            # ✅ 继续下一份文件，别让 session 卡死
            db.session.rollback()

    return results
