import importlib.util
import os
import shutil
import sys
import uuid
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))


def _has_module(name: str) -> bool:
    return importlib.util.find_spec(name) is not None


if not _has_module("flask") or not _has_module("docx"):
    pytest.skip("flask and python-docx are required for KB export tests", allow_module_level=True)

from app import create_app  # noqa: E402
from app.extensions import db  # noqa: E402
from app.models import File, KbBlock, KbDocument  # noqa: E402
from domain.exports.word import WordExportError, export_by_template  # noqa: E402


@pytest.fixture()
def app_ctx(tmp_path):
    os.environ["FLASK_ENV"] = "testing"
    app = create_app("testing")
    app.config["TESTING"] = True

    with app.app_context():
        db.create_all()

    yield app

    with app.app_context():
        db.drop_all()

    storage_dir = Path(app.root_path).parent / "storage"
    if storage_dir.exists():
        shutil.rmtree(storage_dir)


def _create_docx(path: Path, heading: str, body: str):
    from docx import Document

    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    doc.save(str(path))


def _write_template(registry_dir: Path, name: str, content: str):
    registry_dir.mkdir(parents=True, exist_ok=True)
    template_path = registry_dir / name
    template_path.write_text(content, encoding="utf-8")
    return template_path


def test_kb_ingest_creates_blocks(app_ctx):
    app = app_ctx
    client = app.test_client()
    repo_root = Path(app.root_path).parent

    file_id = str(uuid.uuid4())
    rel_path = f"storage/uploads/{file_id}/original.docx"
    abs_path = repo_root / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _create_docx(abs_path, "产品介绍", "这是产品介绍内容")

    with app.app_context():
        db.session.add(
            File(
                id=file_id,
                filename="demo.docx",
                ext="docx",
                size=abs_path.stat().st_size,
                storage_path=rel_path,
            )
        )
        db.session.commit()

    resp = client.post("/api/v1/kb/ingest", json={"file_id": file_id, "title": "投标文件"})
    assert resp.status_code == 200
    data = resp.get_json()
    assert data["doc_id"]
    assert data["block_count"] >= 1

    with app.app_context():
        doc = db.session.get(KbDocument, data["doc_id"])
        assert doc is not None
        blocks = KbBlock.query.filter_by(doc_id=data["doc_id"]).all()
        assert blocks
        for block in blocks:
            assert (repo_root / block.block_docx_path).exists()


def test_export_by_template_outputs_docx(app_ctx):
    app = app_ctx
    repo_root = Path(app.root_path).parent
    registry_dir = REPO_ROOT / "domain" / "templates" / "registry"

    file_id = str(uuid.uuid4())
    rel_path = f"storage/uploads/{file_id}/original.docx"
    abs_path = repo_root / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _create_docx(abs_path, "产品介绍", "这是产品介绍内容")

    with app.app_context():
        db.session.add(
            File(
                id=file_id,
                filename="demo.docx",
                ext="docx",
                size=abs_path.stat().st_size,
                storage_path=rel_path,
            )
        )
        db.session.commit()

    client = app.test_client()
    resp = client.post("/api/v1/kb/ingest", json={"file_id": file_id, "title": "投标文件"})
    assert resp.status_code == 200
    doc_id = resp.get_json()["doc_id"]

    with app.app_context():
        block = KbBlock.query.filter_by(doc_id=doc_id).first()
        assert block is not None
        block.tag = "product_intro"
        db.session.commit()

    template_content = """id: export_demo
version: v1
name: 导出模板
sections:
  - title: 产品介绍
    pick:
      by_tag: ["product_intro"]
      fallback_title_keywords: ["产品介绍"]
      top_k: 5
"""
    template_path = _write_template(registry_dir, "export_demo_v1.yaml", template_content)

    with app.app_context():
        result = export_by_template(job_id="job-export-1", template_id="export_demo", version="v1")

    output_path = repo_root / result["docx_path"]
    assert output_path.exists()

    from docx import Document

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text)
    assert "产品介绍" in full_text

    template_path.unlink(missing_ok=True)


def test_export_by_template_missing_block_fails(app_ctx):
    app = app_ctx
    registry_dir = REPO_ROOT / "domain" / "templates" / "registry"

    template_content = """id: export_empty
version: v1
name: 导出模板
sections:
  - title: 产品介绍
    pick:
      by_tag: ["missing_tag"]
      fallback_title_keywords: ["产品介绍"]
      top_k: 1
"""
    template_path = _write_template(registry_dir, "export_empty_v1.yaml", template_content)

    with app.app_context():
        with pytest.raises(WordExportError, match="no blocks found"):
            export_by_template(job_id="job-export-2", template_id="export_empty", version="v1")

    template_path.unlink(missing_ok=True)
